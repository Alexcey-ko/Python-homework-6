"""Модуль для вывода отчета о структуре каталога в формате DOCX.

Содержит класс DocxWriter.
"""

from docx import Document
from docx.shared import Mm

from python_homework_6.report_manager.base.base_writer import BaseWriter


class DocxWriter(BaseWriter):
    """Класс Writer для создания отчета в DOCX формате."""

    def __init__(self, report_path, dir_path):
        """Инициализация объекта класса.

        Args:
            report_path (str): путь к файлу отчету 
            dir_path (str): путь к исследуемому каталогу
        """
        super().__init__(report_path, dir_path)
        self.__word_doc = None
        self.__data_tab = None

    def create_file(self):
        """Создание Word файла с заголовком и таблицей для вывода данных о структуре каталога."""
        # Создание Word документа
        self.__word_doc = Document()
        self.__word_doc.add_heading(f'Отчет о структуре файлов и папок каталога {self._dir_path}', 0)
        
        #Первая секция документа
        first_section = self.__word_doc.sections[0]
        #Узкие поля (в мм)
        first_section.top_margin = Mm(10)
        first_section.bottom_margin = Mm(10)
        first_section.left_margin = Mm(10)
        first_section.right_margin = Mm(10)

        #Создание таблицы для данных отчета
        self.__data_tab = self.__word_doc.add_table(1, 3, 'Light List Accent 1')
        self.__data_tab.autofit = True
        #Заполнение заголовка таблицы
        for row in self.__data_tab.rows:
            for j, cell in enumerate(row.cells):
                match j:
                    case 0:
                        cell.text = 'Имя файла/папки'
                    case 1:
                        cell.text = 'Размер файла'
                    case 2:
                        cell.text = 'Последнее изменение'

    def write_to_file(self, name, size, last_changed):
        """Вывод информации о файле/папке в файл отчета.

        Args:
            name (str): имя файла/папки
            size (str): размера файла в читаемом формате
            last_changed (str): дата последнего изменения в читаемом формате
        """
        new_row = self.__data_tab.add_row()
        for j, cell in enumerate(new_row.cells):
                    match j:
                        case 0:
                            cell.text = str(name)
                        case 1:
                            cell.text = str(size)
                        case 2:
                            cell.text = str(last_changed) 

    def save_file(self):
        """Сохранение файла отчета."""
        #Сохранение Word документа
        print(self._report_path)
        self.__word_doc.save(self._report_path)